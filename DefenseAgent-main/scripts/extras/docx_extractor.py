"""Example: a custom StructuredExtractor for .docx files.

Drop-in copy for SDK users who want their RAG to ingest Word documents
without forking the framework. Walks paragraphs + inline images + tables,
and persists images to `<resources_dir>/<source_hash>/`.

Usage:
    from DefenseAgent.rag import LlamaIndexRAG, StructuredDocExtractor
    from scripts.extras.docx_extractor import DocxExtractor

    extractor = StructuredDocExtractor(profile)
    extractor.register(DocxExtractor(resources_dir=extractor.resources_dir))
    rag = await LlamaIndexRAG.from_profile(profile, extractor=extractor)

Requires `python-docx`:
    pip install python-docx
"""
from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any

from DefenseAgent.rag.extraction import (
    StructuredChunk,
    StructuredResource,
    _table_to_markdown,
)


def _hash_file(path: Path) -> str:
    """Stable short hash of file contents — same scheme as the built-in extractors."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


class DocxExtractor:
    """Parses .docx into one chunk per Heading-1/2 section.

    For each section, walks paragraphs, tables, and inline images in document
    order; tables get rendered to markdown, images get persisted next to the
    other resources and referenced via `<resource_info>` markers.
    """

    def __init__(self, resources_dir: Path) -> None:
        self.resources_dir = Path(resources_dir)

    def supports(self, source: str | Path) -> bool:
        return Path(source).suffix.lower() == ".docx" and Path(source).is_file()

    def extract(self, source: str | Path) -> list[StructuredChunk]:
        try:
            import docx  # type: ignore  # python-docx
        except ImportError as e:
            raise ImportError(
                "DocxExtractor requires `pip install python-docx`"
            ) from e

        path = Path(source).resolve()
        source_hash = _hash_file(path)[:12]
        sub_dir = self.resources_dir / source_hash
        sub_dir.mkdir(parents=True, exist_ok=True)

        document = docx.Document(str(path))

        # Walk body elements in document order so images / tables / paragraphs
        # interleave correctly. python-docx exposes `document.element.body`
        # which iterates the underlying XML in order.
        chunks: list[StructuredChunk] = []
        current_text: list[str] = []
        current_resources: list[StructuredResource] = []
        section_idx = 0

        def flush_section() -> None:
            nonlocal section_idx
            if not current_text and not current_resources:
                return
            chunks.append(
                StructuredChunk(
                    text="\n\n".join(current_text),
                    resources=list(current_resources),
                    metadata={"source": str(path), "section": section_idx},
                )
            )
            current_text.clear()
            current_resources.clear()
            section_idx += 1

        for block in self._iter_blocks(document):
            kind = block["kind"]
            if kind == "heading":
                # Headings start a new section.
                flush_section()
                current_text.append(str(block["text"]))
            elif kind == "paragraph":
                text = block["text"]
                if text:
                    current_text.append(text)
                # python-docx exposes images via paragraph.runs[*].element ...
                for img_blob in block.get("images", []):
                    res = self._persist_image(
                        img_blob, section_idx, len(current_resources),
                        path, source_hash, sub_dir,
                    )
                    if res is not None:
                        current_resources.append(res)
                        current_text.append(f"<resource_info>{res.id}</resource_info>")
            elif kind == "table":
                md = _table_to_markdown(block["rows"])
                if md:
                    current_text.append(md)
        flush_section()
        return chunks

    @staticmethod
    def _iter_blocks(document: Any) -> Any:
        """Yield paragraphs (with inline images) + tables in document order."""
        from docx.oxml.ns import qn  # type: ignore

        body = document.element.body
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                paragraph = next(
                    (p for p in document.paragraphs if p._element is child),
                    None,
                )
                if paragraph is None:
                    continue
                images: list[bytes] = []
                for run in paragraph.runs:
                    for blip in run.element.findall(f".//{qn('a:blip')}"):
                        rid = blip.get(qn("r:embed"))
                        if rid is None:
                            continue
                        try:
                            part = document.part.related_parts[rid]
                            images.append(part.blob)
                        except Exception:  # noqa: BLE001
                            pass
                style = (paragraph.style.name or "").lower()
                kind = "heading" if style.startswith("heading") else "paragraph"
                yield {"kind": kind, "text": paragraph.text.strip(), "images": images}
            elif tag == "tbl":
                table = next(
                    (t for t in document.tables if t._element is child),
                    None,
                )
                if table is None:
                    continue
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                yield {"kind": "table", "rows": rows}

    def _persist_image(
        self,
        blob: bytes,
        section_idx: int,
        img_idx: int,
        source_path: Path,
        source_hash: str,
        sub_dir: Path,
    ) -> StructuredResource | None:
        # python-docx gives us raw bytes; sniff the format from the header.
        ext = "png" if blob.startswith(b"\x89PNG") else "jpg"
        rid = f"{source_path.name}@{source_hash}@s{section_idx}_img{img_idx}"
        dst = sub_dir / f"s{section_idx}_img{img_idx}.{ext}"
        dst.write_bytes(blob)
        return StructuredResource(
            id=rid,
            kind="image",
            path=dst,
            caption="",  # docx inline images rarely carry alt text in the runs we read
            mime_type=f"image/{ext}",
        )
